import docx
import os

def extract_text_from_txt(txt_file_path: str) -> str:
    """
    Extracts text from a .txt file.

    Args:
        txt_file_path: The file path of the .txt file.

    Returns:
        The extracted text, or an error message if an error occurs.
    """
    try:
        with open(txt_file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except FileNotFoundError:
        return f"Error processing TXT file: File not found at {txt_file_path}"
    except Exception as e:
        return f"Error processing TXT file: {e}"

def extract_text_from_docx(docx_file_path: str) -> str:
    """
    Extracts text from a .docx file.

    Args:
        docx_file_path: The file path of the .docx file.

    Returns:
        The extracted text, or an error message if an error occurs.
    """
    try:
        document = docx.Document(docx_file_path)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except FileNotFoundError:
        return f"Error processing DOCX file: File not found at {docx_file_path}"
    except docx.opc.exceptions.PackageNotFoundError:
        return f"Error processing DOCX file: Not a valid DOCX file or file is corrupted at {docx_file_path}"
    except Exception as e:
        return f"Error processing DOCX file: {e}"

if __name__ == "__main__":
    # Test for extract_text_from_txt
    dummy_txt_path = "dummy_document.txt"
    sample_text_content = "This is a sample text document.\nIt has multiple lines."

    try:
        with open(dummy_txt_path, 'w', encoding='utf-8') as f:
            f.write(sample_text_content)
        print(f"Created dummy TXT file: {dummy_txt_path}")

        extracted_txt = extract_text_from_txt(dummy_txt_path)
        print(f"Extracted text from TXT:\n---\n{extracted_txt}\n---\n")
    except Exception as e:
        print(f"Error in TXT test case setup: {e}")

    # Test error handling for TXT with non-existent file
    non_existent_txt_path = "non_existent_document.txt"
    extracted_text_error_txt = extract_text_from_txt(non_existent_txt_path)
    print(f"Testing TXT with non-existent file: '{extracted_text_error_txt}'\n")

    # Test for extract_text_from_docx
    dummy_docx_path = "dummy_document.docx"

    try:
        doc = docx.Document()
        doc.add_paragraph("This is the first paragraph in the DOCX.")
        doc.add_paragraph("This is the second paragraph, with some more text.")
        doc.add_heading("Heading 1", level=1)
        doc.add_paragraph("Paragraph after heading.")
        doc.save(dummy_docx_path)
        print(f"Created dummy DOCX file: {dummy_docx_path}")

        extracted_docx = extract_text_from_docx(dummy_docx_path)
        print(f"Extracted text from DOCX:\n---\n{extracted_docx}\n---\n")
    except ImportError:
        print("python-docx library is not installed. Cannot create or parse dummy DOCX file.")
    except Exception as e:
        print(f"Error in DOCX test case setup or extraction: {e}")

    # Test error handling for DOCX with non-existent file
    non_existent_docx_path = "non_existent_document.docx"
    extracted_text_error_docx = extract_text_from_docx(non_existent_docx_path)
    print(f"Testing DOCX with non-existent file: '{extracted_text_error_docx}'\n")

    # Test error handling for DOCX with a non-docx file (e.g., the txt file)
    print(f"Testing DOCX extraction with a non-docx file ({dummy_txt_path}):")
    extracted_text_error_invalid_docx = extract_text_from_docx(dummy_txt_path)
    print(f"Result: '{extracted_text_error_invalid_docx}'\n")

    # Clean up dummy files
    if os.path.exists(dummy_txt_path):
        os.remove(dummy_txt_path)
        print(f"Cleaned up {dummy_txt_path}")
    if os.path.exists(dummy_docx_path):
        os.remove(dummy_docx_path)
        print(f"Cleaned up {dummy_docx_path}")
